"""
Text Extraction from Multiple Document Formats
Extracts text content and formatting from various document types.

Supported formats:
- .pages (Apple Pages) - XML-based extraction with PDF preview fallback
- .docx (Microsoft Word) - Full paragraph and table extraction
- .pdf (PDF documents) - Text extraction using pdfplumber
- .txt, .md (Plain text) - Direct text reading with encoding detection

For .pages files:
- Old format: XML-based extraction from index.xml
- Newer format: Falls back to embedded Preview.pdf
- If both fail: Provides manual conversion instructions
"""

import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass, asdict
import hashlib
import tempfile
import os


@dataclass
class FormattingSpan:
    """Represents a span of formatted text."""
    start: int
    end: int
    bold: bool = False
    italic: bool = False


@dataclass
class BulletPoint:
    """Represents a bullet point."""
    text: str
    level: int = 0  # Indentation level


@dataclass
class ExtractionResult:
    """Result of text extraction from a .pages file."""
    text: str
    success: bool
    extraction_method: str  # 'xml', 'pdf_preview', 'failed'
    formatting: Dict = None
    metadata: Dict = None
    error: Optional[str] = None

    def __post_init__(self):
        if self.formatting is None:
            self.formatting = {
                'bold_spans': [],
                'bullets': []
            }
        if self.metadata is None:
            self.metadata = {}

    def to_dict(self) -> Dict:
        """Convert to dictionary for JSON serialization."""
        return asdict(self)


SUPPORTED_EXTENSIONS = {
    '.pages': 'pages',
    '.pdf': 'pdf',
    '.docx': 'docx',
    '.txt': 'text',
    '.md': 'text'
}


def extract_text_from_document(filepath: str) -> Dict:
    """
    Extract text from a document file (any supported format).

    Supports: .pages, .pdf, .docx, .txt, .md

    Args:
        filepath: Path to document file

    Returns:
        Dictionary with extraction results:
        {
            'text': str,
            'success': bool,
            'extraction_method': str,
            'formatting': {...},
            'metadata': {...},
            'error': str (if failed)
        }

    Examples:
        >>> result = extract_text_from_document("document.pages")
        >>> if result['success']:
        ...     print(result['text'])
    """
    path = Path(filepath)

    # Check file exists
    if not path.exists():
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='failed',
            error=f"File not found: {filepath}"
        ).to_dict()

    # Check file extension
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='failed',
            error=f"Unsupported file format: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
        ).to_dict()

    # Route to appropriate extractor
    file_type = SUPPORTED_EXTENSIONS[ext]

    try:
        if file_type == 'pages':
            return _extract_pages(filepath)
        elif file_type == 'pdf':
            return _extract_pdf(filepath)
        elif file_type == 'docx':
            return _extract_docx(filepath)
        elif file_type == 'text':
            return _extract_text(filepath)
    except Exception as e:
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='failed',
            error=f"Extraction failed: {str(e)}"
        ).to_dict()


def _extract_pages(filepath: str) -> Dict:
    """
    Extract text and formatting from a .pages file.

    Args:
        filepath: Path to .pages file

    Returns:
        Dictionary with extraction results
    """
    # .pages files are zip archives
    if not zipfile.is_zipfile(filepath):
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='failed',
            error=f"Invalid .pages file (not a zip archive): {filepath}"
        ).to_dict()

    # Try XML-based extraction (old format)
    result = _try_xml_extraction(filepath)
    if result.success:
        return result.to_dict()

    # Try PDF preview extraction (fallback)
    result = _try_pdf_preview_extraction(filepath)
    if result.success:
        return result.to_dict()

    # Both methods failed
    return ExtractionResult(
        text="",
        success=False,
        extraction_method='failed',
        error=(
            f"Could not extract text from {filepath}. "
            "This .pages file uses a newer format that requires manual conversion. "
            "Please export to .docx or .pdf from Pages and re-run extraction."
        )
    ).to_dict()


def _try_xml_extraction(filepath: str) -> ExtractionResult:
    """
    Try to extract text from XML-based .pages file (old format).

    Args:
        filepath: Path to .pages file

    Returns:
        ExtractionResult
    """
    try:
        with zipfile.ZipFile(filepath, 'r') as zip_ref:
            # Check for index.xml (old format indicator)
            if 'index.xml' not in zip_ref.namelist():
                return ExtractionResult(
                    text="",
                    success=False,
                    extraction_method='xml',
                    error="No index.xml found (not old XML format)"
                )

            # Extract and parse index.xml
            with zip_ref.open('index.xml') as xml_file:
                tree = ET.parse(xml_file)
                root = tree.getroot()

                # Extract text content
                # Pages XML structure varies, but text is typically in <p> or <t> elements
                text_parts = []
                bullets = []

                # Find all text elements
                # Track which elements we've already processed to avoid duplicates
                processed_bullets = set()

                for elem in root.iter():
                    # List items - check multiple possible tag names
                    if (elem.tag.endswith('}list-item') or
                        elem.tag == 'list-item' or
                        elem.tag == 'li'):
                        bullet_text = ''.join(elem.itertext())
                        if bullet_text.strip() and bullet_text not in processed_bullets:
                            processed_bullets.add(bullet_text)
                            bullets.append(BulletPoint(
                                text=bullet_text.strip(),
                                level=0
                            ))
                            text_parts.append(f"• {bullet_text.strip()}")

                    # Text elements (paragraphs)
                    elif elem.tag.endswith('}p') or elem.tag == 'p':
                        para_text = ''.join(elem.itertext())
                        # Skip if this text is already in bullets
                        if para_text.strip() and para_text not in processed_bullets:
                            text_parts.append(para_text.strip())

                if not text_parts:
                    return ExtractionResult(
                        text="",
                        success=False,
                        extraction_method='xml',
                        error="No text content found in XML"
                    )

                full_text = '\n\n'.join(text_parts)

                # Calculate file hash
                file_hash = _calculate_hash(filepath)

                return ExtractionResult(
                    text=full_text,
                    success=True,
                    extraction_method='xml',
                    formatting={
                        'bold_spans': [],  # Could be enhanced
                        'bullets': [asdict(b) for b in bullets]
                    },
                    metadata={
                        'filename': Path(filepath).name,
                        'file_hash': file_hash,
                        'extraction_date': str(__import__('datetime').datetime.now())
                    }
                )

    except Exception as e:
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='xml',
            error=f"XML extraction failed: {str(e)}"
        )


def _try_pdf_preview_extraction(filepath: str) -> ExtractionResult:
    """
    Try to extract text from embedded PDF preview.

    Newer .pages files contain QuickLook/Preview.pdf which can be extracted.
    Requires pdfplumber to be installed.

    Args:
        filepath: Path to .pages file

    Returns:
        ExtractionResult
    """
    try:
        import pdfplumber
    except ImportError:
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='pdf_preview',
            error="pdfplumber not installed (required for PDF preview extraction)"
        )

    try:
        with zipfile.ZipFile(filepath, 'r') as zip_ref:
            # Look for Preview.pdf
            preview_paths = [
                'QuickLook/Preview.pdf',
                'preview.pdf',
                'Preview.pdf'
            ]

            pdf_path = None
            for path in preview_paths:
                if path in zip_ref.namelist():
                    pdf_path = path
                    break

            if not pdf_path:
                return ExtractionResult(
                    text="",
                    success=False,
                    extraction_method='pdf_preview',
                    error="No Preview.pdf found in .pages archive"
                )

            # Extract PDF to temporary file
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_pdf = os.path.join(temp_dir, 'preview.pdf')
                with open(temp_pdf, 'wb') as f:
                    f.write(zip_ref.read(pdf_path))

                # Extract text from PDF
                text_parts = []
                with pdfplumber.open(temp_pdf) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)

                if not text_parts:
                    return ExtractionResult(
                        text="",
                        success=False,
                        extraction_method='pdf_preview',
                        error="No text extracted from PDF preview"
                    )

                full_text = '\n\n'.join(text_parts)
                file_hash = _calculate_hash(filepath)

                return ExtractionResult(
                    text=full_text,
                    success=True,
                    extraction_method='pdf_preview',
                    formatting={
                        'bold_spans': [],  # PDF extraction doesn't preserve formatting well
                        'bullets': []
                    },
                    metadata={
                        'filename': Path(filepath).name,
                        'file_hash': file_hash,
                        'extraction_date': str(__import__('datetime').datetime.now()),
                        'note': 'Extracted from PDF preview - formatting may not be preserved'
                    }
                )

    except Exception as e:
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='pdf_preview',
            error=f"PDF preview extraction failed: {str(e)}"
        )


def _extract_pdf(filepath: str) -> Dict:
    """
    Extract text from PDF file.

    Args:
        filepath: Path to PDF file

    Returns:
        Dictionary with extraction results
    """
    try:
        import pdfplumber
    except ImportError:
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='pdf',
            error="pdfplumber not installed. Install with: pip install pdfplumber"
        ).to_dict()

    try:
        text_parts = []
        metadata = {
            'file_type': 'pdf',
            'filename': Path(filepath).name,
            'page_count': 0
        }

        with pdfplumber.open(filepath) as pdf:
            metadata['page_count'] = len(pdf.pages)

            # Extract PDF metadata if available
            if pdf.metadata:
                metadata.update({
                    'pdf_title': pdf.metadata.get('Title', ''),
                    'pdf_author': pdf.metadata.get('Author', ''),
                    'pdf_subject': pdf.metadata.get('Subject', ''),
                    'pdf_created': pdf.metadata.get('CreationDate', '')
                })

            # Extract text from each page
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        if not text_parts:
            return ExtractionResult(
                text="",
                success=False,
                extraction_method='pdf',
                error="No text content found in PDF"
            ).to_dict()

        full_text = '\n\n'.join(text_parts)
        metadata['file_hash'] = _calculate_hash(filepath)
        metadata['extraction_date'] = str(__import__('datetime').datetime.now())

        return ExtractionResult(
            text=full_text,
            success=True,
            extraction_method='pdf',
            metadata=metadata
        ).to_dict()

    except Exception as e:
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='pdf',
            error=f"PDF extraction failed: {str(e)}"
        ).to_dict()


def _extract_docx(filepath: str) -> Dict:
    """
    Extract text from Word document (.docx).

    Args:
        filepath: Path to .docx file

    Returns:
        Dictionary with extraction results
    """
    try:
        from docx import Document
    except ImportError:
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='docx',
            error="python-docx not installed. Install with: pip install python-docx"
        ).to_dict()

    try:
        text_parts = []
        metadata = {
            'file_type': 'docx',
            'filename': Path(filepath).name,
            'paragraph_count': 0,
            'table_count': 0
        }

        # Load document
        doc = Document(filepath)

        # Extract core properties if available
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata['doc_title'] = core_props.title
            if core_props.author:
                metadata['doc_author'] = core_props.author
            if core_props.created:
                metadata['doc_created'] = str(core_props.created)
            if core_props.modified:
                metadata['doc_modified'] = str(core_props.modified)
        except Exception as e:
            # Core properties may not be available in some documents
            pass

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        metadata['paragraph_count'] = len(doc.paragraphs)

        # Extract tables
        if doc.tables:
            metadata['table_count'] = len(doc.tables)
            for table in doc.tables:
                # Extract table as tab-separated text
                for row in table.rows:
                    row_text = '\t'.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

        if not text_parts:
            return ExtractionResult(
                text="",
                success=False,
                extraction_method='docx',
                error="No text content found in document"
            ).to_dict()

        full_text = '\n'.join(text_parts)
        metadata['file_hash'] = _calculate_hash(filepath)
        metadata['extraction_date'] = str(__import__('datetime').datetime.now())

        return ExtractionResult(
            text=full_text,
            success=True,
            extraction_method='docx',
            metadata=metadata
        ).to_dict()

    except Exception as e:
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='docx',
            error=f"DOCX extraction failed: {str(e)}"
        ).to_dict()


def _extract_text(filepath: str) -> Dict:
    """
    Extract text from plain text file (.txt, .md).

    Args:
        filepath: Path to text file

    Returns:
        Dictionary with extraction results
    """
    try:
        metadata = {
            'file_type': 'text',
            'filename': Path(filepath).name
        }

        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    text = f.read()

                metadata['encoding'] = encoding
                metadata['char_count'] = len(text)
                metadata['line_count'] = text.count('\n') + 1
                metadata['file_hash'] = _calculate_hash(filepath)
                metadata['extraction_date'] = str(__import__('datetime').datetime.now())

                return ExtractionResult(
                    text=text,
                    success=True,
                    extraction_method='text',
                    metadata=metadata
                ).to_dict()
            except UnicodeDecodeError:
                continue

        # If all encodings fail
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='text',
            error="Unable to decode text file with supported encodings (utf-8, latin-1, cp1252)"
        ).to_dict()

    except Exception as e:
        return ExtractionResult(
            text="",
            success=False,
            extraction_method='text',
            error=f"Text extraction failed: {str(e)}"
        ).to_dict()


def extract_metadata(filepath: str, text: str) -> Dict:
    """
    Extract metadata from filename and text content.

    Identifies:
    - Target position/organization from text
    - Document classification hints

    Args:
        filepath: Path to document
        text: Extracted text content

    Returns:
        Dictionary with metadata
    """
    from utils.date_parser import extract_date_from_filename

    filename = Path(filepath).name

    metadata = {
        'filename': filename,
        'filepath': str(filepath),
        'date': None,
        'target_position': None,
        'target_organization': None
    }

    # Extract date from filename
    doc_date = extract_date_from_filename(filename)
    if doc_date:
        metadata['date'] = doc_date.isoformat()

    # Try to extract target position/organization from filename
    # Common patterns: "2024-03-15-cover-letter-company-position.pages"
    parts = filename.lower().replace('.pages', '').split('-')

    # Look for organization hints in filename
    for i, part in enumerate(parts):
        if part in ['cover', 'letter', 'resume', 'cv']:
            # Next parts might be organization/position
            if i + 1 < len(parts):
                metadata['target_organization'] = ' '.join(parts[i+1:]).title()
                break

    # Try to extract from text content
    if text:
        # Look for "Dear [Name]" or "Dear Hiring Manager at [Company]"
        import re
        # Try pattern: "Dear Hiring Manager at Company"
        dear_match = re.search(r'Dear\s+(?:Hiring Manager|Search Committee)\s+(?:at|for)\s+([A-Z][A-Za-z\s&\.]+?)(?:,|\n)', text)
        if dear_match:
            org = dear_match.group(1).strip()
            if not metadata['target_organization']:
                metadata['target_organization'] = org
        else:
            # Try simpler pattern: "Dear [Name/Title],"
            dear_match = re.search(r'Dear\s+([A-Z][^,\n]+),', text)
            if dear_match:
                potential_org = dear_match.group(1).strip()
                # Only use if it looks like an organization (contains certain keywords or is reasonably short)
                if any(keyword in potential_org.lower() for keyword in ['hiring', 'search', 'committee', 'team']):
                    pass  # This is likely a title, not organization
                elif not metadata['target_organization'] and len(potential_org) < 50:
                    # Could be a person's name or short organization
                    pass  # Skip for now unless we have better heuristics

        # Look for position titles in first paragraph
        position_keywords = ['position', 'role', 'opportunity', 'director', 'manager', 'professor', 'dean']
        first_para = text[:500]  # Check first 500 chars
        for keyword in position_keywords:
            pattern = rf'([A-Z][^.\n]*{keyword}[^.\n]*)'
            match = re.search(pattern, first_para, re.IGNORECASE)
            if match:
                position = match.group(1).strip()
                if len(position) < 100:  # Reasonable length
                    metadata['target_position'] = position
                    break

    return metadata


def _calculate_hash(filepath: str) -> str:
    """Calculate SHA-256 hash of file."""
    sha256_hash = hashlib.sha256()
    with open(filepath, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()


def get_extraction_instructions() -> str:
    """
    Get instructions for manual extraction when automated extraction fails.

    Returns:
        Instruction string for user
    """
    return """
Automated extraction from .pages file failed.

To manually convert your .pages file:

1. Open the document in Apple Pages
2. File > Export To > choose format:
   - Word (.docx) - Recommended, preserves formatting
   - PDF - Good for final documents
3. Save the exported file in the same directory
4. Re-run the lexicon builder with the exported file

Tip: You can batch export multiple .pages files:
- Select files in Finder
- Right-click > Quick Actions > Convert with Pages
- Choose output format

The lexicon builder supports: .docx, .pdf, .txt
"""
