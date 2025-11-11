"""Apply styles to document content using template."""
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from typing import List, Dict, Any, Optional
from pathlib import Path
from lxml import etree
import logging

from cv_formatting.play_titles_lookup import PlayTitlesLookup

logger = logging.getLogger(__name__)


class StyleApplicator:
    """Applies styles to content using template document."""

    def __init__(self, template_path: str,
                 dictionary_path: Optional[str] = None,
                 signature_path: Optional[str] = None):
        """
        Initialize with template.

        Args:
            template_path: Path to cv-template.docx
            dictionary_path: Path to play-titles-dictionary.yaml (optional)
            signature_path: Path to signatures directory (optional)
        """
        self.template_path = template_path
        self.document_type = 'cv'  # Default to CV mode
        self.play_lookup = PlayTitlesLookup(dictionary_path) if dictionary_path else None
        self.signature_path = Path(signature_path) if signature_path else None

    def apply_styles(self, content_mapping: List[Dict[str, Any]],
                     output_path: str, document_type: str = 'cv',
                     metadata: Optional[Dict[str, Any]] = None) -> bool:
        """
        Apply styles to content and save document.

        Args:
            content_mapping: List of content elements with styles
            output_path: Where to save formatted document
            document_type: Type of document ('cv' or 'cover-letter')
            metadata: Optional document metadata (for headers, etc.)

        Returns:
            True if successful

        Content mapping format:
            [{
                "text": "Content text",
                "style": "Style Name",
                "type": "paragraph" | "inline" | "image",
                "inline_styles": [  # optional, for inline styling
                    {"text": "text", "style": "Style Name"}
                ],
                "runs": [  # optional, legacy format
                    {"text": "text", "style": "Style Name"}
                ]
            }]
        """
        try:
            # Store document type for conditional formatting
            self.document_type = document_type
            metadata = metadata or {}

            # Load template
            doc = Document(self.template_path)

            # Add page headers if specified in metadata
            if metadata.get('page_header', {}).get('enabled'):
                self._add_page_headers(doc, metadata)

            # Track first section for spacing adjustments
            self.first_section_seen = False

            # Apply each content element
            for i, item in enumerate(content_mapping):
                prev_item = content_mapping[i-1] if i > 0 else None
                next_item = content_mapping[i+1] if i < len(content_mapping)-1 else None
                self._add_content_item(doc, item, prev_item, next_item)

            # Save formatted document
            doc.save(output_path)
            logger.info(f"Document saved: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to apply styles: {e}")
            return False

    def _add_content_item(self, doc: Document, item: Dict[str, Any],
                          prev_item: Optional[Dict[str, Any]] = None,
                          next_item: Optional[Dict[str, Any]] = None):
        """Add single content item to document."""
        item_type = item.get('type', 'paragraph')

        if item_type == 'image':
            # Handle signature image
            self._add_signature_image(doc, item)
            return

        if item_type == 'paragraph':
            # Add paragraph with style
            style_name = item.get('style', 'Body Text')

            # Check if item has inline styling (new format)
            if 'inline_styles' in item or (self.play_lookup and style_name == 'Body Text'):
                para = doc.add_paragraph(style=style_name)
                self._apply_inline_styles(para, item)

            # Legacy format: inline runs
            elif 'runs' in item:
                para = doc.add_paragraph(style=style_name)
                for run_data in item['runs']:
                    run = para.add_run(run_data['text'])
                    if run_data.get('style'):
                        run.style = run_data['style']

            # Simple paragraph
            else:
                para = doc.add_paragraph(item['text'], style=style_name)

            # Apply context-aware formatting for Section Header or RE Line
            if style_name == 'Section Header':
                is_first_section = not self.first_section_seen
                self._apply_section_header_formatting(para, is_first_section)
                self.first_section_seen = True
            elif style_name == 'RE Line':
                self._apply_re_line_formatting(para)

            # Normal spacing - overlap handled by floating signature positioning
            if next_item and next_item.get('type') == 'image':
                # Paragraph before signature (e.g., "Sincerely,")
                para.paragraph_format.space_after = Pt(0)
            elif prev_item and prev_item.get('type') == 'image':
                # Paragraph after signature (e.g., "Anthony Byrnes")
                para.paragraph_format.space_before = Pt(0)

    def _apply_section_header_formatting(self, para, is_first_section=False):
        """
        Apply context-aware formatting to Section Header paragraph.

        Section headers are ALWAYS orange for both CVs and cover letters.
        CV mode: Orange, 11pt, Bold
        Cover letter mode: Orange, 13pt, Bold

        Args:
            para: The paragraph to format
            is_first_section: True if this is the first section header in the document
        """
        # Always orange for both document types
        color = RGBColor(255, 109, 73)

        if self.document_type == 'cv':
            size = Pt(11)
        else:  # cover-letter
            size = Pt(13)

        # Apply direct formatting to all runs in the paragraph
        for run in para.runs:
            run.font.color.rgb = color
            run.font.size = size
            run.font.bold = True

        # First section should have no space before
        if is_first_section:
            para.paragraph_format.space_before = Pt(0)

    def _apply_re_line_formatting(self, para):
        """
        Apply formatting to RE Line paragraph.

        RE Line: Orange, 13pt, Bold (same as cover letter section headers)
        """
        color = RGBColor(255, 109, 73)  # Orange
        size = Pt(13)

        for run in para.runs:
            run.font.color.rgb = color
            run.font.size = size
            run.font.bold = True

    def _apply_inline_styles(self, para, item: Dict[str, Any]):
        """
        Apply inline styles to paragraph text using hybrid approach.

        Combines:
        1. Dictionary lookup for known plays (auto-styled)
        2. Explicit inline_styles from JSON (overrides)
        """
        text = item.get('text', '')
        explicit_styles = item.get('inline_styles', [])

        # Find plays from dictionary (if enabled)
        auto_plays = []
        if self.play_lookup and self.document_type == 'cover-letter':
            auto_plays = self.play_lookup.find_plays_in_text(text)

        # Build combined style map: {(start, end): style_name}
        style_map = {}

        # Add dictionary plays (unless excluded)
        for play_text, start, end in auto_plays:
            if not self.play_lookup.should_exclude(play_text, explicit_styles):
                style_map[(start, end)] = 'Play Title'

        # Add explicit inline styles (override dictionary)
        for style_spec in explicit_styles:
            if not style_spec.get('exclude'):
                style_text = style_spec['text']
                style_name = style_spec.get('style')

                # Find all occurrences
                start_pos = 0
                while True:
                    pos = text.find(style_text, start_pos)
                    if pos == -1:
                        break
                    style_map[(pos, pos + len(style_text))] = style_name
                    start_pos = pos + len(style_text)

        # Apply styles by building runs
        if not style_map:
            # No inline styles, just add plain text
            para.add_run(text)
            return

        # Sort style regions by position
        regions = sorted(style_map.items())

        # Build runs
        current_pos = 0
        for (start, end), style_name in regions:
            # Add text before this styled region
            if start > current_pos:
                para.add_run(text[current_pos:start])

            # Add styled text
            run = para.add_run(text[start:end])
            if style_name:
                run.style = style_name

            current_pos = end

        # Add remaining text
        if current_pos < len(text):
            para.add_run(text[current_pos:])

    def _add_signature_image(self, doc: Document, item: Dict[str, Any]):
        """
        Add signature image to document.

        Looks for image in signatures/ directory.
        """
        if not self.signature_path:
            logger.warning("Signature path not configured, skipping image")
            return

        # Get signature name from item (defaults to "signature")
        sig_name = item.get('text', 'signature')
        image_file = self.signature_path / f"{sig_name}.png"

        if not image_file.exists():
            logger.warning(f"Signature image not found: {image_file}")
            return

        # Add signature as inline image with minimal spacing and height
        # Note: Floating images with overlap cause document corruption
        # Using inline with tight spacing and smaller height instead
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run()
        try:
            # Use smaller width and let height auto-scale
            # This reduces the vertical space taken by the signature
            run.add_picture(str(image_file), width=Inches(1.2))
            logger.info(f"Added signature image: {image_file}")
        except Exception as e:
            logger.error(f"Failed to add signature image: {e}")

    def _add_page_headers(self, doc: Document, metadata: Dict[str, Any]):
        """
        Add page headers to document.

        Headers start from page 2 (not page 1).
        Format matches Colburn style: gray text, 0.25" left indent, lowercase "page X"
        """
        page_header = metadata.get('page_header', {})
        left_text = page_header.get('left', '')
        right_text = page_header.get('right', '')

        if not left_text and not right_text:
            return

        # Enable different first page header
        section = doc.sections[0]
        section.different_first_page_header_footer = True

        # First page header stays empty (already is by default)

        # Add header for pages 2+
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        # Use XML to set up paragraph with proper tab stops and formatting
        # This matches the Colburn style exactly
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Set up paragraph properties
        pPr = header_para._element.get_or_add_pPr()

        # Add tab stops (right-aligned tab at 9360 twips = 6.5 inches from left)
        # No left indent - header should align with body text
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9360')  # 6.5 inches
        tabs.append(tab)
        pPr.append(tabs)

        # Add left text (author name + organization)
        # Format: "ANTHONY BYRNES - Title Case Position"
        if left_text:
            run = header_para.add_run(left_text)
            run.bold = True
            run.font.name = 'Helvetica'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)  # Gray

        # Add single tab to push page number to the right
        if left_text and right_text:
            run = header_para.add_run('\t')
            run.font.name = 'Helvetica'

        # Add right text ("page") with page number field
        if right_text:
            # Add "page " text (lowercase, 10pt, bold, gray, Helvetica)
            page_run = header_para.add_run(right_text + ' ')
            page_run.bold = True
            page_run.font.name = 'Helvetica'
            page_run.font.size = Pt(10)
            page_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray

            # Build complete field XML with all runs
            # Note: Page number may display as {PAGE} until Word updates fields.
            # Word will auto-update when document is opened/printed/saved.
            # User can manually update with F9 or right-click > Update Field.
            field_begin = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:rPr>'
                '<w:rFonts w:ascii="Helvetica" w:hAnsi="Helvetica"/>'
                '<w:b/>'
                '<w:sz w:val="20"/>'
                '<w:color w:val="808080"/>'
                '</w:rPr>'
                '<w:fldChar w:fldCharType="begin"/>'
                '</w:r>'
            )

            field_instr = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:rPr>'
                '<w:rFonts w:ascii="Helvetica" w:hAnsi="Helvetica"/>'
                '<w:b/>'
                '<w:sz w:val="20"/>'
                '<w:color w:val="808080"/>'
                '</w:rPr>'
                '<w:instrText xml:space="preserve"> PAGE </w:instrText>'
                '</w:r>'
            )

            field_sep = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:rPr>'
                '<w:rFonts w:ascii="Helvetica" w:hAnsi="Helvetica"/>'
                '<w:b/>'
                '<w:sz w:val="20"/>'
                '<w:color w:val="808080"/>'
                '</w:rPr>'
                '<w:fldChar w:fldCharType="separate"/>'
                '</w:r>'
            )

            field_end = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:rPr>'
                '<w:rFonts w:ascii="Helvetica" w:hAnsi="Helvetica"/>'
                '<w:b/>'
                '<w:sz w:val="20"/>'
                '<w:color w:val="808080"/>'
                '</w:rPr>'
                '<w:fldChar w:fldCharType="end"/>'
                '</w:r>'
            )

            # Append all field components
            header_para._element.append(field_begin)
            header_para._element.append(field_instr)
            header_para._element.append(field_sep)
            header_para._element.append(field_end)

        logger.info(f"Added page headers (Colburn style): '{left_text}' | '{right_text}'")
