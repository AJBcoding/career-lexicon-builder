"""Text extraction from various document formats.

Supports PDF, Word (.docx), and plain text files.
"""

from typing import Dict, Optional, Tuple
from pathlib import Path
import pdfplumber
from docx import Document
from dataclasses import dataclass
import hashlib


@dataclass
class ExtractionResult:
    """Result of text extraction from a document."""
    text: str
    metadata: Dict
    success: bool
    error: Optional[str] = None


class TextExtractor:
    """Extract text from various document formats."""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.txt': 'text',
        '.md': 'text'
    }
    
    def __init__(self):
        """Initialize text extractor."""
        pass
    
    def extract(self, filepath: str) -> ExtractionResult:
        """Extract text from a document file.
        
        Args:
            filepath: Path to document file
            
        Returns:
            ExtractionResult containing text, metadata, and status
        """
        path = Path(filepath)
        
        # Check file exists
        if not path.exists():
            return ExtractionResult(
                text="",
                metadata={},
                success=False,
                error=f"File not found: {filepath}"
            )
        
        # Get file extension
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return ExtractionResult(
                text="",
                metadata={},
                success=False,
                error=f"Unsupported file format: {ext}"
            )
        
        # Extract based on file type
        file_type = self.SUPPORTED_EXTENSIONS[ext]
        
        try:
            if file_type == 'pdf':
                return self._extract_pdf(path)
            elif file_type == 'docx':
                return self._extract_docx(path)
            elif file_type == 'text':
                return self._extract_text(path)
        except Exception as e:
            return ExtractionResult(
                text="",
                metadata={},
                success=False,
                error=f"Extraction failed: {str(e)}"
            )
    
    def _extract_pdf(self, path: Path) -> ExtractionResult:
        """Extract text from PDF file.
        
        Args:
            path: Path to PDF file
            
        Returns:
            ExtractionResult with extracted text
        """
        text_parts = []
        metadata = {
            'file_type': 'pdf',
            'filename': path.name,
            'page_count': 0
        }
        
        with pdfplumber.open(path) as pdf:
            metadata['page_count'] = len(pdf.pages)
            
            # Extract PDF metadata if available
            if pdf.metadata:
                metadata.update({
                    'pdf_title': pdf.metadata.get('Title', ''),
                    'pdf_author': pdf.metadata.get('Author', ''),
                    'pdf_subject': pdf.metadata.get('Subject', ''),
                    'pdf_created': pdf.metadata.get('CreationDate', '')
                })
            
            # Extract text from each page
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        full_text = '\n\n'.join(text_parts)
        
        # Add file hash
        metadata['file_hash'] = self._calculate_hash(path)
        
        return ExtractionResult(
            text=full_text,
            metadata=metadata,
            success=True
        )
    
    def _extract_docx(self, path: Path) -> ExtractionResult:
        """Extract text from Word document.
        
        Args:
            path: Path to .docx file
            
        Returns:
            ExtractionResult with extracted text
        """
        text_parts = []
        metadata = {
            'file_type': 'docx',
            'filename': path.name,
            'paragraph_count': 0,
            'table_count': 0
        }
        
        # Load document
        doc = Document(path)
        
        # Extract core properties if available
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata['doc_title'] = core_props.title
            if core_props.author:
                metadata['doc_author'] = core_props.author
            if core_props.created:
                metadata['doc_created'] = str(core_props.created)
            if core_props.modified:
                metadata['doc_modified'] = str(core_props.modified)
        except:
            pass  # Core properties may not be available
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        metadata['paragraph_count'] = len(doc.paragraphs)
        
        # Extract tables
        if doc.tables:
            metadata['table_count'] = len(doc.tables)
            for table in doc.tables:
                # Extract table as tab-separated text
                for row in table.rows:
                    row_text = '\t'.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
        
        full_text = '\n'.join(text_parts)
        
        # Add file hash
        metadata['file_hash'] = self._calculate_hash(path)
        
        return ExtractionResult(
            text=full_text,
            metadata=metadata,
            success=True
        )
    
    def _extract_text(self, path: Path) -> ExtractionResult:
        """Extract text from plain text file.
        
        Args:
            path: Path to text file
            
        Returns:
            ExtractionResult with extracted text
        """
        metadata = {
            'file_type': 'text',
            'filename': path.name
        }
        
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    text = f.read()
                
                metadata['encoding'] = encoding
                metadata['char_count'] = len(text)
                metadata['line_count'] = text.count('\n') + 1
                metadata['file_hash'] = self._calculate_hash(path)
                
                return ExtractionResult(
                    text=text,
                    metadata=metadata,
                    success=True
                )
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail
        return ExtractionResult(
            text="",
            metadata=metadata,
            success=False,
            error="Unable to decode text file with supported encodings"
        )
    
    def _calculate_hash(self, path: Path) -> str:
        """Calculate SHA-256 hash of file.
        
        Args:
            path: Path to file
            
        Returns:
            Hex string of file hash
        """
        sha256_hash = hashlib.sha256()
        
        with open(path, "rb") as f:
            # Read file in chunks
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        
        return sha256_hash.hexdigest()
    
    def get_supported_extensions(self) -> list:
        """Get list of supported file extensions.
        
        Returns:
            List of supported extensions (with dots)
        """
        return list(self.SUPPORTED_EXTENSIONS.keys())


def extract_text(filepath: str) -> Tuple[str, Dict]:
    """Convenience function to extract text from a file.
    
    Args:
        filepath: Path to document file
        
    Returns:
        Tuple of (text, metadata)
        
    Raises:
        ValueError: If extraction fails
    """
    extractor = TextExtractor()
    result = extractor.extract(filepath)
    
    if not result.success:
        raise ValueError(f"Text extraction failed: {result.error}")
    
    return result.text, result.metadata
