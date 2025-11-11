import pytest
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from cv_formatting.template_builder import TemplateBuilder
from cv_formatting.style_mapping import get_all_semantic_styles


def test_create_template_generates_docx(tmp_path):
    """Test that template creation generates .docx file"""
    output_path = tmp_path / "test-template.docx"

    builder = TemplateBuilder()
    result = builder.create_template(str(output_path))

    assert result is True
    assert output_path.exists()


def test_template_contains_all_semantic_styles(tmp_path):
    """Test that template includes all 12 semantic styles"""
    output_path = tmp_path / "test-template.docx"

    builder = TemplateBuilder()
    builder.create_template(str(output_path))

    # Load and verify
    doc = Document(str(output_path))
    style_names = {s.name for s in doc.styles}

    expected_styles = get_all_semantic_styles()

    for expected in expected_styles:
        assert expected in style_names, f"Missing style: {expected}"


def test_template_styles_have_correct_properties(tmp_path):
    """Test that styles have expected formatting properties"""
    output_path = tmp_path / "test-template.docx"

    builder = TemplateBuilder()
    builder.create_template(str(output_path))

    doc = Document(str(output_path))

    # Check Section Header style
    section_header = doc.styles['Section Header']
    assert section_header.font.bold == True
    assert section_header.font.size.pt == 10
    assert section_header.font.color.rgb is not None  # Has orange color

    # Check Body Text style
    body_text = doc.styles['Body Text']
    assert body_text.font.size.pt == 9
    assert body_text.font.name == 'Helvetica'


def test_template_contains_date_line_style(tmp_path):
    """Test that template includes Date Line style."""
    output_path = tmp_path / "test-template.docx"

    builder = TemplateBuilder()
    builder.create_template(str(output_path))

    doc = Document(str(output_path))

    style_names = [s.name for s in doc.styles]
    assert 'Date Line' in style_names, "Date Line style not found in template"

    # Verify Date Line properties
    date_line_style = doc.styles['Date Line']
    assert date_line_style.type == WD_STYLE_TYPE.PARAGRAPH

    # Check alignment
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    assert date_line_style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_template_contains_13_styles(tmp_path):
    """Test that template contains all 13 expected styles."""
    expected_styles = [
        'CV Name', 'Section Header', 'Body Text', 'Timeline Entry',
        'Bullet Standard', 'Bullet Gray', 'Bullet Emphasis',
        'Play Title', 'Institution', 'Job Title', 'Orange Emphasis', 'Gray Text',
        'Date Line'  # NEW
    ]

    output_path = tmp_path / "test-template.docx"

    builder = TemplateBuilder()
    builder.create_template(str(output_path))

    doc = Document(str(output_path))

    style_names = [s.name for s in doc.styles]

    for expected in expected_styles:
        assert expected in style_names, f"Style '{expected}' not found in template"
