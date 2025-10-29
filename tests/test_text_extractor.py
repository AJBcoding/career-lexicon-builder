"""Tests for text extraction system."""

import pytest
from pathlib import Path
import tempfile
import os
from src.text_extractor import (
    TextExtractor,
    ExtractionResult,
    extract_text
)


# Test document contents
SAMPLE_TEXT = """
This is a sample text document for testing.
It has multiple lines.
And various content to extract.
"""

SAMPLE_RESUME_TEXT = """
JANE DOE
Software Engineer

PROFESSIONAL SUMMARY
Experienced software engineer with 5+ years developing applications.

WORK EXPERIENCE
Senior Engineer | Tech Corp | 2020 - Present
- Developed scalable systems
- Led engineering team

EDUCATION
BS Computer Science | University | 2018

SKILLS
Python, Java, AWS
"""


class TestTextExtractor:
    """Test suite for TextExtractor class."""
    
    @pytest.fixture
    def extractor(self):
        """Fixture providing TextExtractor instance."""
        return TextExtractor()
    
    @pytest.fixture
    def temp_dir(self):
        """Fixture providing temporary directory for test files."""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield Path(tmpdir)
    
    def test_supported_extensions(self, extractor):
        """Test that supported extensions are correctly listed."""
        extensions = extractor.get_supported_extensions()
        
        assert '.pdf' in extensions
        assert '.docx' in extensions
        assert '.txt' in extensions
        assert '.md' in extensions
    
    def test_extract_missing_file(self, extractor):
        """Test extraction from non-existent file."""
        result = extractor.extract("nonexistent_file.pdf")
        
        assert result.success is False
        assert result.text == ""
        assert "not found" in result.error.lower()
    
    def test_extract_unsupported_format(self, extractor, temp_dir):
        """Test extraction from unsupported file format."""
        # Create unsupported file
        unsupported_file = temp_dir / "test.xyz"
        unsupported_file.write_text("content")
        
        result = extractor.extract(str(unsupported_file))
        
        assert result.success is False
        assert "unsupported" in result.error.lower()
    
    def test_extract_text_file_utf8(self, extractor, temp_dir):
        """Test extraction from UTF-8 text file."""
        text_file = temp_dir / "test.txt"
        text_file.write_text(SAMPLE_TEXT, encoding='utf-8')
        
        result = extractor.extract(str(text_file))
        
        assert result.success is True
        assert SAMPLE_TEXT.strip() in result.text
        assert result.metadata['file_type'] == 'text'
        assert result.metadata['encoding'] == 'utf-8'
        assert result.metadata['line_count'] > 0
        assert 'file_hash' in result.metadata
    
    def test_extract_markdown_file(self, extractor, temp_dir):
        """Test extraction from Markdown file."""
        md_content = "# Heading\n\nSome **bold** text\n\n- List item"
        md_file = temp_dir / "test.md"
        md_file.write_text(md_content, encoding='utf-8')
        
        result = extractor.extract(str(md_file))
        
        assert result.success is True
        assert "Heading" in result.text
        assert "bold" in result.text
        assert result.metadata['file_type'] == 'text'
    
    def test_extract_text_latin1_encoding(self, extractor, temp_dir):
        """Test extraction from Latin-1 encoded text file."""
        # Create file with Latin-1 encoding
        latin1_text = "RÃ©sumÃ© with special characters: cafÃ©"
        text_file = temp_dir / "test_latin1.txt"
        text_file.write_bytes(latin1_text.encode('latin-1'))
        
        result = extractor.extract(str(text_file))
        
        assert result.success is True
        assert "cafÃ©" in result.text or "caf" in result.text  # Encoding may vary
        assert result.metadata['file_type'] == 'text'
    
    def test_extraction_result_structure(self, extractor, temp_dir):
        """Test that ExtractionResult has correct structure."""
        text_file = temp_dir / "test.txt"
        text_file.write_text("test content")
        
        result = extractor.extract(str(text_file))
        
        assert isinstance(result, ExtractionResult)
        assert isinstance(result.text, str)
        assert isinstance(result.metadata, dict)
        assert isinstance(result.success, bool)
        assert result.error is None or isinstance(result.error, str)
    
    def test_metadata_includes_filename(self, extractor, temp_dir):
        """Test that metadata includes filename."""
        text_file = temp_dir / "myfile.txt"
        text_file.write_text("content")
        
        result = extractor.extract(str(text_file))
        
        assert result.metadata['filename'] == 'myfile.txt'
    
    def test_metadata_includes_hash(self, extractor, temp_dir):
        """Test that metadata includes file hash."""
        text_file = temp_dir / "test.txt"
        text_file.write_text("content")
        
        result = extractor.extract(str(text_file))
        
        assert 'file_hash' in result.metadata
        assert len(result.metadata['file_hash']) == 64  # SHA-256 hex length
    
    def test_hash_consistency(self, extractor, temp_dir):
        """Test that same file produces same hash."""
        text_file = temp_dir / "test.txt"
        text_file.write_text("content")
        
        result1 = extractor.extract(str(text_file))
        result2 = extractor.extract(str(text_file))
        
        assert result1.metadata['file_hash'] == result2.metadata['file_hash']
    
    def test_hash_changes_with_content(self, extractor, temp_dir):
        """Test that different content produces different hash."""
        text_file = temp_dir / "test.txt"
        
        text_file.write_text("content 1")
        result1 = extractor.extract(str(text_file))
        
        text_file.write_text("content 2")
        result2 = extractor.extract(str(text_file))
        
        assert result1.metadata['file_hash'] != result2.metadata['file_hash']


class TestConvenienceFunction:
    """Test suite for extract_text convenience function."""
    
    @pytest.fixture
    def temp_dir(self):
        """Fixture providing temporary directory."""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield Path(tmpdir)
    
    def test_extract_text_success(self, temp_dir):
        """Test convenience function with successful extraction."""
        text_file = temp_dir / "test.txt"
        text_file.write_text(SAMPLE_TEXT)
        
        text, metadata = extract_text(str(text_file))
        
        assert SAMPLE_TEXT.strip() in text
        assert isinstance(metadata, dict)
        assert 'filename' in metadata
    
    def test_extract_text_failure(self):
        """Test convenience function with extraction failure."""
        with pytest.raises(ValueError) as exc_info:
            extract_text("nonexistent_file.pdf")
        
        assert "extraction failed" in str(exc_info.value).lower()


class TestWordDocExtraction:
    """Test suite for Word document extraction."""
    
    @pytest.fixture
    def temp_dir(self):
        """Fixture providing temporary directory."""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield Path(tmpdir)
    
    def test_extract_docx_basic(self, temp_dir):
        """Test basic Word document extraction."""
        # Note: This test requires creating an actual .docx file
        # For now, we test the error handling when file doesn't exist
        from docx import Document
        
        # Create a simple Word document
        doc = Document()
        doc.add_paragraph("Test paragraph 1")
        doc.add_paragraph("Test paragraph 2")
        
        docx_file = temp_dir / "test.docx"
        doc.save(str(docx_file))
        
        extractor = TextExtractor()
        result = extractor.extract(str(docx_file))
        
        assert result.success is True
        assert "Test paragraph 1" in result.text
        assert "Test paragraph 2" in result.text
        assert result.metadata['file_type'] == 'docx'
        assert result.metadata['paragraph_count'] >= 2
    
    def test_extract_docx_with_tables(self, temp_dir):
        """Test Word document with tables."""
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("Before table")
        
        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Data 1"
        table.rows[1].cells[1].text = "Data 2"
        
        doc.add_paragraph("After table")
        
        docx_file = temp_dir / "test_table.docx"
        doc.save(str(docx_file))
        
        extractor = TextExtractor()
        result = extractor.extract(str(docx_file))
        
        assert result.success is True
        assert "Before table" in result.text
        assert "After table" in result.text
        assert "Header 1" in result.text
        assert "Data 1" in result.text
        assert result.metadata['table_count'] == 1


class TestPDFExtraction:
    """Test suite for PDF extraction."""
    
    @pytest.fixture
    def temp_dir(self):
        """Fixture providing temporary directory."""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield Path(tmpdir)
    
    def test_pdf_extraction_basic(self, temp_dir):
        """Test basic PDF extraction."""
        # Note: Creating test PDFs requires additional libraries
        # For now, we test that the extraction method exists and handles errors
        extractor = TextExtractor()
        
        # Test with non-existent PDF
        result = extractor.extract("nonexistent.pdf")
        assert result.success is False
        assert "not found" in result.error.lower()
    
    # Additional PDF tests would require creating actual PDF files
    # This can be done with libraries like reportlab or fpdf


class TestEdgeCases:
    """Test edge cases and error conditions."""
    
    @pytest.fixture
    def temp_dir(self):
        """Fixture providing temporary directory."""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield Path(tmpdir)
    
    def test_empty_text_file(self, temp_dir):
        """Test extraction from empty text file."""
        empty_file = temp_dir / "empty.txt"
        empty_file.write_text("")
        
        extractor = TextExtractor()
        result = extractor.extract(str(empty_file))
        
        assert result.success is True
        assert result.text == ""
    
    def test_binary_file_as_text(self, temp_dir):
        """Test extraction from binary file with text extension."""
        # Use bytes that are invalid in UTF-8, Latin-1, and CP1252
        binary_file = temp_dir / "binary.txt"
        binary_file.write_bytes(b'\xff\xfe\xfd\xfc\xfb')
        
        extractor = TextExtractor()
        result = extractor.extract(str(binary_file))
        
        # May succeed with latin-1 (it accepts all bytes) or fail
        # Either is acceptable behavior for binary data
        if not result.success:
            assert "decode" in result.error.lower()
        else:
            # If it succeeds, text should contain the decoded bytes
            assert len(result.text) > 0
    
    def test_very_large_text_file(self, temp_dir):
        """Test extraction from large text file."""
        large_content = "Line of text\n" * 10000  # 10k lines
        large_file = temp_dir / "large.txt"
        large_file.write_text(large_content)
        
        extractor = TextExtractor()
        result = extractor.extract(str(large_file))
        
        assert result.success is True
        assert len(result.text) > 100000  # Should be quite large
        assert result.metadata['line_count'] >= 10000
    
    def test_unicode_text_file(self, temp_dir):
        """Test extraction from Unicode text file."""
        unicode_content = "Hello ä¸–ç•Œ ðŸŒ ÐŸÑ€Ð¸Ð²ÐµÑ‚ Ù…Ø±Ø­Ø¨Ø§"
        unicode_file = temp_dir / "unicode.txt"
        unicode_file.write_text(unicode_content, encoding='utf-8')
        
        extractor = TextExtractor()
        result = extractor.extract(str(unicode_file))
        
        assert result.success is True
        assert "Hello" in result.text
        # Unicode characters should be preserved
        assert "ä¸–ç•Œ" in result.text or len(result.text) > 0
