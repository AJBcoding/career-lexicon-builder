"""Tests for CV page headers."""
import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt
from cv_formatting.style_applicator import StyleApplicator


def test_cv_page_headers_enabled(tmp_path):
    """Test that page headers appear on page 2+ when enabled for CV"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-cv-with-headers.docx"

    # Create multi-page CV content
    content_mapping = [
        {"text": "ANTHONY BYRNES", "style": "CV Name", "type": "paragraph"},
        {"text": "T: 213.305.3132", "style": "Contact Info", "type": "paragraph"},
        {"text": "E: anthonybyrnes@mac.com", "style": "Contact Info", "type": "paragraph"},
        {"text": "EDUCATION", "style": "Section Header", "type": "paragraph"},
    ]

    # Add enough content to force page 2
    for i in range(20):
        content_mapping.append({
            "text": f"2020-2024  California State University, Long Beach - Position {i+1}",
            "style": "Timeline Entry",
            "type": "paragraph"
        })
        content_mapping.append({
            "text": f"This is a detailed description of the role with enough text to fill space. " * 15,
            "style": "Body Text",
            "type": "paragraph"
        })

    # Add metadata with CV page header configuration
    metadata = {
        "type": "cv",
        "author_name": "Anthony Byrnes",
        "document_title": "Curriculum Vitae",
        "page_header": {
            "enabled": True,
            "left": "ANTHONY BYRNES - Curriculum Vitae",
            "right": "page"
        }
    }

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(
        content_mapping,
        str(output),
        document_type='cv',
        metadata=metadata
    )

    assert result is True
    assert output.exists()

    # Verify headers exist
    doc = Document(str(output))
    section = doc.sections[0]

    # Verify different first page is enabled
    assert section.different_first_page_header_footer is True, \
        "Should have different first page header"

    # Verify first page header is empty
    first_page_header = section.first_page_header
    first_page_text = "".join([p.text for p in first_page_header.paragraphs]).strip()
    assert first_page_text == "", f"First page header should be empty, got: {first_page_text}"

    # Verify subsequent pages header contains CV name and title
    header = section.header
    header_text = "".join([p.text for p in header.paragraphs])
    assert "ANTHONY BYRNES" in header_text, \
        f"Header should contain 'ANTHONY BYRNES', got: {header_text}"
    assert "Curriculum Vitae" in header_text, \
        f"Header should contain 'Curriculum Vitae', got: {header_text}"


def test_cv_page_headers_disabled(tmp_path):
    """Test that page headers do not appear when disabled for CV"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-cv-no-headers.docx"

    content_mapping = [
        {"text": "ANTHONY BYRNES", "style": "CV Name", "type": "paragraph"},
        {"text": "EDUCATION", "style": "Section Header", "type": "paragraph"},
    ]

    metadata = {
        "type": "cv",
        "author_name": "Anthony Byrnes",
        "page_header": {
            "enabled": False
        }
    }

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(
        content_mapping,
        str(output),
        document_type='cv',
        metadata=metadata
    )

    assert result is True
    assert output.exists()

    # Verify no headers
    doc = Document(str(output))
    section = doc.sections[0]

    # Either different_first_page is False, or headers are empty
    if section.different_first_page_header_footer:
        header = section.header
        header_text = "".join([p.text for p in header.paragraphs]).strip()
        assert header_text == "", f"Header should be empty when disabled, got: {header_text}"


def test_cv_page_headers_colburn_style(tmp_path):
    """Test that CV page headers match Colburn formatting (gray, Helvetica, no indent, lowercase 'page')"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-cv-colburn-headers.docx"

    # Create multi-page CV content
    content_mapping = [
        {"text": "ANTHONY BYRNES", "style": "CV Name", "type": "paragraph"},
        {"text": "T: 213.305.3132", "style": "Contact Info", "type": "paragraph"},
        {"text": "EDUCATION", "style": "Section Header", "type": "paragraph"},
    ]

    # Add enough content to force page 2
    for i in range(20):
        content_mapping.append({
            "text": f"This is paragraph {i+1} with enough text to fill multiple pages. " * 30,
            "style": "Body Text",
            "type": "paragraph"
        })

    # Add metadata with page header configuration
    metadata = {
        "type": "cv",
        "author_name": "Anthony Byrnes",
        "document_title": "Curriculum Vitae",
        "page_header": {
            "enabled": True,
            "left": "ANTHONY BYRNES - Curriculum Vitae",
            "right": "page"  # lowercase
        }
    }

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(
        content_mapping,
        str(output),
        document_type='cv',
        metadata=metadata
    )

    assert result is True
    assert output.exists()

    # Verify header formatting matches Colburn style
    doc = Document(str(output))
    section = doc.sections[0]

    # Verify different first page is enabled
    assert section.different_first_page_header_footer is True

    # Verify header exists and contains expected text
    header = section.header
    header_text = "".join([p.text for p in header.paragraphs])
    assert "ANTHONY BYRNES - Curriculum Vitae" in header_text
    assert "page" in header_text.lower()

    # Check formatting on header paragraph
    header_para = header.paragraphs[0]

    # Check no left indent (should align with body text)
    left_indent = header_para.paragraph_format.left_indent
    assert left_indent is None or left_indent.twips == 0, \
        f"Expected no left indent (0 twips), got {left_indent.twips if left_indent else 'None'}"

    # Check that runs have correct formatting
    gray_found = False
    helvetica_found = False
    bold_found = False
    correct_size_found = False

    for run in header_para.runs:
        # Check gray color (808080 = RGB 128,128,128)
        if run.font.color.rgb and run.font.color.rgb == (128, 128, 128):
            gray_found = True

        # Check Helvetica font
        if run.font.name and "Helvetica" in run.font.name:
            helvetica_found = True

        # Check bold
        if run.bold:
            bold_found = True

        # Check 10pt size
        if run.font.size and run.font.size.pt == 10:
            correct_size_found = True

    assert gray_found, "Expected at least one run to have gray color (RGB 128,128,128)"
    assert helvetica_found, "Expected Helvetica font to be used"
    assert bold_found, "Expected bold text in header"
    assert correct_size_found, "Expected 10pt font size in header"


def test_cv_page_headers_format_variation(tmp_path):
    """Test different CV page header format variations"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-cv-header-variations.docx"

    content_mapping = [
        {"text": "ANTHONY BYRNES", "style": "CV Name", "type": "paragraph"},
        {"text": "EDUCATION", "style": "Section Header", "type": "paragraph"},
    ]

    # Add content to force multiple pages
    for i in range(20):
        content_mapping.append({
            "text": f"Content line {i+1}. " * 40,
            "style": "Body Text",
            "type": "paragraph"
        })

    # Test with just "CV" as document title
    metadata = {
        "type": "cv",
        "author_name": "Anthony Byrnes",
        "document_title": "CV",
        "page_header": {
            "enabled": True,
            "left": "ANTHONY BYRNES - CV",
            "right": "page"
        }
    }

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(
        content_mapping,
        str(output),
        document_type='cv',
        metadata=metadata
    )

    assert result is True
    assert output.exists()

    doc = Document(str(output))
    section = doc.sections[0]
    header = section.header
    header_text = "".join([p.text for p in header.paragraphs])

    # Verify short format works
    assert "ANTHONY BYRNES - CV" in header_text
    assert "page" in header_text.lower()


def test_cv_section_header_size(tmp_path):
    """Test that CV section headers are 11pt (not 13pt like cover letters)"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-cv-section-header-size.docx"

    content_mapping = [
        {"text": "EDUCATION", "style": "Section Header", "type": "paragraph"},
        {"text": "Some content", "style": "Body Text", "type": "paragraph"},
    ]

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(
        content_mapping,
        str(output),
        document_type='cv',
        metadata={}
    )

    assert result is True
    assert output.exists()

    doc = Document(str(output))

    # Find the section header paragraph
    section_header_para = None
    for para in doc.paragraphs:
        if para.text == "EDUCATION":
            section_header_para = para
            break

    assert section_header_para is not None, "Could not find EDUCATION section header"

    # Check that section header is 11pt for CV
    for run in section_header_para.runs:
        if run.text:
            assert run.font.size.pt == 11, \
                f"CV section headers should be 11pt, got {run.font.size.pt}pt"
            break
