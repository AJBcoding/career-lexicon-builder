def test_ucla_cao_letter_formatting():
    """Integration test: Format actual UCLA CAO cover letter."""
    from pathlib import Path
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from cv_formatting.style_applicator import StyleApplicator
    import tempfile

    template = Path("cv_formatting/templates/career-documents-template.docx")
    fixture = Path(__file__).parent / 'fixtures' / 'ucla-cao-cover-letter-v3.md'

    # For this test, we'll create a simplified content mapping
    # In real usage, Claude skill would parse the markdown and create this mapping
    content_mapping = [
        # Date (from actual letter - though not in the fixture, would be added)
        # {"text": "November 25, 2024", "style": "Date Line", "type": "paragraph"},
        # {"text": "", "style": "Body Text", "type": "paragraph"},

        # Salutation
        {"text": "Dear Members of the UCLA School of Theater, Film and Television Search Committee,",
         "style": "Body Text", "type": "paragraph"},
        {"text": "", "style": "Body Text", "type": "paragraph"},

        # First paragraph (would have inline Institution, Job Title styles in real usage)
        {"text": "I am writing to express my interest in the Associate Dean and Chief Administrative Officer position...",
         "style": "Body Text", "type": "paragraph"},
        {"text": "", "style": "Body Text", "type": "paragraph"},

        # Section header (thematic)
        {"text": "Why UCLA? Why TFT? Why now?", "style": "Section Header", "type": "paragraph"},
        {"text": "", "style": "Body Text", "type": "paragraph"},

        # Body paragraph
        {"text": "The university's commitment to inclusive excellence aligns with my values.",
         "style": "Body Text", "type": "paragraph"},
        {"text": "", "style": "Body Text", "type": "paragraph"},

        # Closing
        {"text": "Thank you for your time and attention,", "style": "Body Text", "type": "paragraph"},
        {"text": "", "style": "Body Text", "type": "paragraph"},
        {"text": "Anthony Byrnes", "style": "Body Text", "type": "paragraph"},
    ]

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        output_path = Path(tmp.name)

    try:
        applicator = StyleApplicator(str(template))
        result = applicator.apply_styles(content_mapping, str(output_path), document_type='cover-letter')

        assert result is True, "UCLA CAO letter formatting failed"
        assert output_path.exists(), "Output file not created"

        # Verify document structure
        doc = Document(str(output_path))
        assert len(doc.paragraphs) > 0, "Document has no paragraphs"

        # Find and verify section header formatting
        section_header_found = False
        for para in doc.paragraphs:
            if "Why UCLA" in para.text:
                section_header_found = True
                run = para.runs[0] if para.runs else None
                if run:
                    # Check: Black, 13pt, Bold
                    assert run.font.color.rgb == RGBColor(0, 0, 0), \
                        f"Section header should be black, got {run.font.color.rgb}"
                    assert run.font.size == Pt(13), \
                        f"Section header should be 13pt, got {run.font.size}"
                    assert run.font.bold is True, "Section header should be bold"
                break

        assert section_header_found, "Section header 'Why UCLA?' not found in document"

    finally:
        if output_path.exists():
            output_path.unlink()
