import pytest
from pathlib import Path
from docx import Document
from cv_formatting.style_applicator import StyleApplicator


def test_apply_paragraph_style(tmp_path):
    """Test applying paragraph style to content"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-output.docx"

    content_mapping = [
        {
            "text": "EDUCATION",
            "style": "Section Header",
            "type": "paragraph"
        },
        {
            "text": "Standard body paragraph text.",
            "style": "Body Text",
            "type": "paragraph"
        }
    ]

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(content_mapping, str(output))

    assert result is True
    assert output.exists()

    # Verify styles applied
    doc = Document(str(output))
    assert len(doc.paragraphs) == 2
    assert doc.paragraphs[0].style.name == "Section Header"
    assert doc.paragraphs[1].style.name == "Body Text"


def test_apply_character_style_inline(tmp_path):
    """Test applying character style within paragraph"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-output.docx"

    content_mapping = [
        {
            "text": "2023 - Present ",
            "style": "Timeline Entry",
            "type": "paragraph",
            "runs": [
                {"text": "2023 - Present ", "style": None},
                {"text": "California State University", "style": "Institution"}
            ]
        }
    ]

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(content_mapping, str(output))

    assert result is True

    # Verify inline style applied
    doc = Document(str(output))
    para = doc.paragraphs[0]
    assert para.style.name == "Timeline Entry"
    assert len(para.runs) == 2
    assert para.runs[1].style.name == "Institution"


def test_section_header_cv_mode(tmp_path):
    """Test Section Header formatting in CV mode (orange, 11pt)"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-cv-section-header.docx"

    content_mapping = [
        {
            "text": "EDUCATION",
            "style": "Section Header",
            "type": "paragraph"
        }
    ]

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(content_mapping, str(output), document_type='cv')

    assert result is True
    assert output.exists()

    # Verify Section Header has CV formatting (orange, 11pt, bold)
    doc = Document(str(output))
    para = doc.paragraphs[0]
    assert para.style.name == "Section Header"

    # Check run formatting (direct formatting overrides style)
    run = para.runs[0]
    assert run.font.size.pt == 11, f"Expected 11pt, got {run.font.size.pt}pt"
    assert run.font.bold is True, "Expected bold"

    # Check color is orange (#FF6D49 = RGB(255, 109, 73))
    assert run.font.color.rgb is not None, "Expected color to be set"
    rgb = run.font.color.rgb
    assert rgb == (255, 109, 73), f"Expected RGB(255, 109, 73), got RGB{rgb}"


def test_section_header_cover_letter_mode(tmp_path):
    """Test Section Header formatting in cover letter mode (black, 13pt)"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-cover-letter-section-header.docx"

    content_mapping = [
        {
            "text": "INTRODUCTION",
            "style": "Section Header",
            "type": "paragraph"
        }
    ]

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(content_mapping, str(output), document_type='cover-letter')

    assert result is True
    assert output.exists()

    # Verify Section Header has cover letter formatting (black, 13pt, bold)
    doc = Document(str(output))
    para = doc.paragraphs[0]
    assert para.style.name == "Section Header"

    # Check run formatting (direct formatting overrides style)
    run = para.runs[0]
    assert run.font.size.pt == 13, f"Expected 13pt, got {run.font.size.pt}pt"
    assert run.font.bold is True, "Expected bold"

    # Check color is black (RGB(0, 0, 0))
    assert run.font.color.rgb is not None, "Expected color to be set"
    rgb = run.font.color.rgb
    assert rgb == (0, 0, 0), f"Expected RGB(0, 0, 0), got RGB{rgb}"
