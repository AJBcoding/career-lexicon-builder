import pytest
from pathlib import Path
from docx import Document
from cv_formatting.style_applicator import StyleApplicator


def test_cover_letter_page_headers_enabled(tmp_path):
    """Test that page headers appear on page 2+ when enabled"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-cover-letter-with-headers.docx"

    # Create multi-page content
    content_mapping = [
        {"text": "ANTHONY BYRNES", "style": "Contact Name", "type": "paragraph"},
        {"text": "T: 213.305.3132", "style": "Contact Info", "type": "paragraph"},
        {"text": "E: anthonybyrnes@mac.com", "style": "Contact Info", "type": "paragraph"},
        {"text": "November 11, 2025", "style": "Date Line", "type": "paragraph"},
        {"text": "UCLA School of Theater, Film and Television", "style": "Recipient Address", "type": "paragraph"},
        {"text": "RE: Associate Dean and Chief Administrative Officer [CAO]", "style": "RE Line", "type": "paragraph"},
        {"text": "Dear Members of the Search Committee,", "style": "Body Text", "type": "paragraph"},
    ]

    # Add enough body paragraphs to force page 2
    for i in range(15):
        content_mapping.append({
            "text": f"This is paragraph {i+1} with enough text to fill multiple pages. " * 20,
            "style": "Body Text",
            "type": "paragraph"
        })

    # Add metadata with page header configuration
    metadata = {
        "type": "cover-letter",
        "author_name": "Anthony Byrnes",
        "page_header": {
            "enabled": True,
            "left": "ANTHONY BYRNES",
            "right": "Page"
        }
    }

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(
        content_mapping,
        str(output),
        document_type='cover-letter',
        metadata=metadata
    )

    assert result is True
    assert output.exists()

    # Verify headers exist
    doc = Document(str(output))

    # Check that document has multiple sections or a header
    assert len(doc.sections) > 0, "Document should have sections"

    # Get the primary header (for pages 2+)
    # Note: python-docx represents headers per section
    # For a document with different first page header, we need to check:
    # - section.different_first_page_header_footer should be True
    # - section.first_page_header should be empty or minimal
    # - section.header should contain our header text

    section = doc.sections[0]

    # Verify different first page is enabled
    assert section.different_first_page_header_footer is True, \
        "Should have different first page header"

    # Verify first page header is empty or minimal
    first_page_header = section.first_page_header
    first_page_text = "".join([p.text for p in first_page_header.paragraphs]).strip()
    assert first_page_text == "", f"First page header should be empty, got: {first_page_text}"

    # Verify subsequent pages header contains author name
    header = section.header
    header_text = "".join([p.text for p in header.paragraphs])
    assert "ANTHONY BYRNES" in header_text, \
        f"Header should contain 'ANTHONY BYRNES', got: {header_text}"


def test_cover_letter_page_headers_disabled(tmp_path):
    """Test that page headers do not appear when disabled"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-cover-letter-no-headers.docx"

    content_mapping = [
        {"text": "ANTHONY BYRNES", "style": "Contact Name", "type": "paragraph"},
        {"text": "Body text here.", "style": "Body Text", "type": "paragraph"},
    ]

    metadata = {
        "type": "cover-letter",
        "author_name": "Anthony Byrnes",
        "page_header": {
            "enabled": False
        }
    }

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(
        content_mapping,
        str(output),
        document_type='cover-letter',
        metadata=metadata
    )

    assert result is True
    assert output.exists()

    # Verify no headers
    doc = Document(str(output))
    section = doc.sections[0]

    # Either different_first_page is False, or headers are empty
    if section.different_first_page_header_footer:
        header = section.header
        header_text = "".join([p.text for p in header.paragraphs]).strip()
        assert header_text == "", f"Header should be empty when disabled, got: {header_text}"


def test_cover_letter_page_headers_colburn_style(tmp_path):
    """Test that page headers match Colburn formatting (gray, 0.25in indent, lowercase 'page')"""
    template = Path("cv_formatting/templates/career-documents-template.docx")
    output = tmp_path / "test-cover-letter-colburn-headers.docx"

    # Create multi-page content
    content_mapping = [
        {"text": "ANTHONY BYRNES", "style": "Contact Name", "type": "paragraph"},
        {"text": "T: 213.305.3132", "style": "Contact Info", "type": "paragraph"},
    ]

    # Add enough content to force page 2
    for i in range(15):
        content_mapping.append({
            "text": f"This is paragraph {i+1} with enough text to fill multiple pages. " * 20,
            "style": "Body Text",
            "type": "paragraph"
        })

    # Add metadata with page header configuration including organization
    metadata = {
        "type": "cover-letter",
        "author_name": "Anthony Byrnes",
        "document_title": "UCLA Cover Letter",
        "page_header": {
            "enabled": True,
            "left": "ANTHONY BYRNES - UCLA Cover Letter",
            "right": "page"  # lowercase
        }
    }

    applicator = StyleApplicator(str(template))
    result = applicator.apply_styles(
        content_mapping,
        str(output),
        document_type='cover-letter',
        metadata=metadata
    )

    assert result is True
    assert output.exists()

    # Verify header formatting matches Colburn style
    doc = Document(str(output))
    section = doc.sections[0]

    # Verify different first page is enabled
    assert section.different_first_page_header_footer is True

    # Verify header exists and contains expected text
    header = section.header
    header_text = "".join([p.text for p in header.paragraphs])
    assert "ANTHONY BYRNES - UCLA Cover Letter" in header_text
    assert "page" in header_text.lower()

    # Check formatting on header paragraph
    header_para = header.paragraphs[0]

    # Check no left indent (should align with body text)
    left_indent = header_para.paragraph_format.left_indent
    assert left_indent is None or left_indent.twips == 0, \
        f"Expected no left indent (0 twips), got {left_indent.twips if left_indent else 'None'}"

    # Check that runs have correct formatting
    gray_found = False
    helvetica_found = False

    for run in header_para.runs:
        # Check gray color (808080 = RGB 128,128,128)
        if run.font.color.rgb and run.font.color.rgb == (128, 128, 128):
            gray_found = True

        # Check Helvetica font
        if run.font.name and "Helvetica" in run.font.name:
            helvetica_found = True

        # Check 10pt size for name/title
        if "ANTHONY" in run.text:
            assert run.font.size.pt == 10, f"Expected 10pt for name, got {run.font.size.pt}pt"

    assert gray_found, "Expected at least one run to have gray color (RGB 128,128,128)"
    assert helvetica_found, "Expected Helvetica font to be used"
