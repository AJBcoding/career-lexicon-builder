"""Tests for cover letter formatting functionality."""
import pytest
from pathlib import Path
from docx import Document
from cv_formatting.style_applicator import StyleApplicator


@pytest.fixture
def cover_letter_fixture():
    """Path to sample cover letter markdown fixture."""
    return Path(__file__).parent / "fixtures" / "sample_cover_letter.md"


@pytest.fixture
def template_path():
    """Path to career documents template."""
    return Path("cv_formatting/templates/career-documents-template.docx")


def test_cover_letter_structure_detection(cover_letter_fixture):
    """
    Test that fixture contains expected structural elements.

    Verifies the sample cover letter includes:
    - Date line
    - Salutation
    - Body paragraphs mentioning institutions, productions, positions
    - Thematic section headers (markdown headers)
    - Closing
    - Signature
    """
    assert cover_letter_fixture.exists(), "Fixture file not found"

    # Read fixture content
    content = cover_letter_fixture.read_text()

    # Check for date line (should be at the start)
    assert "December 15, 2024" in content or any(
        line.strip() and not line.startswith('#') and ',' in line
        for line in content.split('\n')[:3]
    ), "Date line not found at document start"

    # Check for salutation
    assert "Dear" in content, "Salutation not found"

    # Check for thematic section headers (markdown headers)
    assert "# Why" in content or "#" in content, "Thematic section headers (markdown) not found"

    # Check for institutions mentioned
    assert "Northwestern" in content or "University" in content, "Institution names not found"

    # Check for productions mentioned (italicized titles)
    assert "*Hamlet*" in content or "*" in content, "Production titles not found"

    # Check for positions/roles mentioned
    assert "Professor" in content or "Director" in content, "Position/role titles not found"

    # Check for closing
    assert "Sincerely" in content or "Regards" in content, "Closing not found"

    # Check for signature (name should appear after closing)
    lines = content.split('\n')
    sincerely_idx = next((i for i, line in enumerate(lines) if 'Sincerely' in line or 'Regards' in line), -1)
    assert sincerely_idx >= 0, "Could not find closing line"

    # Check that there's content after closing (signature)
    remaining_content = '\n'.join(lines[sincerely_idx + 1:]).strip()
    assert len(remaining_content) > 0, "Signature not found after closing"
    assert any(char.isalpha() for char in remaining_content), "Signature should contain name"


def test_cover_letter_formatting_basic(cover_letter_fixture, template_path, tmp_path):
    """
    Test that basic formatting pipeline works with cover-letter document type.

    Verifies:
    - StyleApplicator accepts document_type='cover-letter'
    - Content mapping can be applied to create formatted document
    - Output is valid .docx file
    """
    output = tmp_path / "test-cover-letter.docx"

    # Create simple content mapping representing cover letter structure
    # This simulates what a parser would generate from the markdown
    content_mapping = [
        {
            "text": "December 15, 2024",
            "style": "Date Line",
            "type": "paragraph"
        },
        {
            "text": "Dear Search Committee,",
            "style": "Body Text",
            "type": "paragraph"
        },
        {
            "text": "I am writing to express my interest in the Associate Professor of Theatre position at Northwestern University.",
            "style": "Body Text",
            "type": "paragraph"
        },
        {
            "text": "Why Northwestern University?",
            "style": "Section Header",
            "type": "paragraph"
        },
        {
            "text": "Northwestern's commitment to innovative theatrical practice aligns perfectly with my artistic vision.",
            "style": "Body Text",
            "type": "paragraph"
        },
        {
            "text": "Professional Experience",
            "style": "Section Header",
            "type": "paragraph"
        },
        {
            "text": "Throughout my tenure at University of California, Berkeley, I have directed over twenty major productions.",
            "style": "Body Text",
            "type": "paragraph"
        },
        {
            "text": "Sincerely,",
            "style": "Body Text",
            "type": "paragraph"
        },
        {
            "text": "Dr. Sarah Martinez",
            "style": "Body Text",
            "type": "paragraph"
        }
    ]

    # Apply styles using cover-letter document type
    applicator = StyleApplicator(str(template_path))
    result = applicator.apply_styles(
        content_mapping,
        str(output),
        document_type='cover-letter'
    )

    # Verify successful formatting
    assert result is True, "StyleApplicator.apply_styles() should return True"
    assert output.exists(), "Output .docx file should be created"

    # Verify output is valid .docx
    try:
        doc = Document(str(output))
        assert len(doc.paragraphs) > 0, "Document should contain paragraphs"

        # Verify document contains expected number of paragraphs
        assert len(doc.paragraphs) == len(content_mapping), \
            f"Expected {len(content_mapping)} paragraphs, got {len(doc.paragraphs)}"

        # Verify Section Header paragraphs have cover-letter formatting (black, 13pt, bold)
        section_headers = [p for p in doc.paragraphs if p.style.name == 'Section Header']
        assert len(section_headers) == 2, "Should have 2 section headers"

        for header in section_headers:
            # Check that direct formatting was applied (cover letter mode)
            run = header.runs[0]
            assert run.font.size.pt == 13, \
                f"Section Header in cover letter should be 13pt, got {run.font.size.pt}pt"
            assert run.font.bold is True, "Section Header should be bold"

            # Check color is black (RGB(0, 0, 0))
            assert run.font.color.rgb == (0, 0, 0), \
                f"Section Header in cover letter should be black, got {run.font.color.rgb}"

        # Verify Date Line style is applied correctly
        date_para = doc.paragraphs[0]
        assert date_para.style.name == 'Date Line', "First paragraph should use Date Line style"

    except Exception as e:
        pytest.fail(f"Failed to open/validate output .docx: {e}")
